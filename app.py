import os
import re
import uuid
import io
import time
import threading
import subprocess
import imageio_ffmpeg
from flask import Flask, render_template, request, jsonify, send_file
from openai import OpenAI
import pandas as pd
from docx import Document

app = Flask(__name__)

client = OpenAI(
    api_key=os.environ.get("GROQ_API_KEY"),
    base_url="https://api.groq.com/openai/v1"
)

UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# Memoria para tareas en segundo plano
TAREAS = {}
RESULTS_CACHE = {}

MARKET_KEYWORDS = {
    "Precio/Costos": ["precio", "costo", "caro", "barato", "oferta", "pago", "dinero", "soles", "dólares", "inversión"],
    "Calidad/Producto": ["calidad", "bueno", "malo", "excelente", "falla", "material", "duradero"],
    "Servicio/Atención": ["atención", "servicio", "soporte", "ayuda", "rápido", "lento", "queja"],
    "Cantidad/Ventas": ["cantidad", "vendido", "unidades", "stock", "total", "volumen"]
}

def obtener_ffmpeg():
    """Obtiene la ruta del ejecutable de FFmpeg instalado vía pip"""
    return imageio_ffmpeg.get_ffmpeg_exe()

def generar_resumen_ia(texto):
    if not texto.strip(): return "No hay texto suficiente."
    
    # Truncado Defensivo (Límite de Llama 3)
    limite_caracteres = 18000
    texto_seguro = texto[:limite_caracteres]
    aviso = "\n\n⚠️ NOTA DE HELIOS: Debido a la longitud extrema, este resumen ejecutivo se generó analizando la primera parte de la reunión para respetar los límites de la IA. Tus archivos descargables contienen la transcripción completa." if len(texto) > limite_caracteres else ""
    
    prompt = f"""Actúa como consultor senior. Haz un RESUMEN EJECUTIVO identificando Datos Duros (precios, costos) y Tendencias. Usa viñetas formales.\n\nTranscripción:\n{texto_seguro}"""
    
    try:
        response = client.chat.completions.create(
            model="llama-3.1-8b-instant",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.2
        )
        return response.choices[0].message.content + aviso
    except Exception as e:
        return f"Error al resumir: {str(e)}"

def procesar_audio_pesado(task_id, rutas_archivos):
    """Esta función corre en segundo plano para evitar que Render corte la conexión"""
    try:
        TAREAS[task_id]['estado'] = 'Cortando y preparando audios (FFmpeg)...'
        
        carpeta_trabajo = os.path.join(UPLOAD_FOLDER, task_id)
        os.makedirs(carpeta_trabajo, exist_ok=True)
        
        pedazos_totales = []
        ffmpeg_exe = obtener_ffmpeg()

        # 1. AUTOCORTE DE AUDIOS LARGOS
        for ruta in rutas_archivos:
            nombre_base = os.path.basename(ruta).split('.')[0]
            # Cortar en pedazos de 15 min (900 seg) y comprimir a mp3
            comando = [
                ffmpeg_exe, "-i", ruta, "-f", "segment", "-segment_time", "900",
                "-c:a", "libmp3lame", "-b:a", "64k", 
                os.path.join(carpeta_trabajo, f"{nombre_base}_parte_%03d.mp3")
            ]
            subprocess.run(comando, check=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
            os.remove(ruta) # Borrar el wav original pesado
        
        # Recoger todos los pedazos generados
        pedazos_totales = sorted([os.path.join(carpeta_trabajo, f) for f in os.listdir(carpeta_trabajo) if f.endswith('.mp3')])
        total_archivos = len(pedazos_totales)
        
        transcripciones_list = []
        conteo_acumulado = {cat: 0 for cat in MARKET_KEYWORDS.keys()}

        # 2. TRANSCRIPCIÓN CON AUTO-RETRY
        for idx, pedazo in enumerate(pedazos_totales):
            nombre_archivo = os.path.basename(pedazo)
            TAREAS[task_id]['estado'] = f'Transcribiendo parte {idx+1} de {total_archivos}...'
            
            exito = False
            intentos = 0
            while not exito and intentos < 3:
                try:
                    with open(pedazo, "rb") as audio_file:
                        transcription = client.audio.transcriptions.create(
                            model="whisper-large-v3", file=audio_file, language="es"
                        )
                    texto_final = transcription.text
                    transcripciones_list.append({"archivo": nombre_archivo, "texto": texto_final})
                    
                    texto_limpio = re.sub(r'[^\w\s]', '', texto_final.lower())
                    for cat, sin in MARKET_KEYWORDS.items():
                        for pal in sin:
                            conteo_acumulado[cat] += len(re.findall(r'\b' + re.escape(pal) + r'\b', texto_limpio))
                    exito = True
                    
                except Exception as e:
                    error_str = str(e)
                    if "429" in error_str and "rate_limit_exceeded" in error_str:
                        # Extraer el tiempo de espera del mensaje de Groq
                        match = re.search(r'try again in (?:(\d+)m)?(?:([0-9.]+)s)', error_str)
                        if match:
                            mins = int(match.group(1)) if match.group(1) else 0
                            secs = float(match.group(2)) if match.group(2) else 0
                            espera = (mins * 60) + secs + 5 # 5 segundos extra de margen
                            TAREAS[task_id]['estado'] = f'Límite de Groq alcanzado. Helios pausado automáticamente por {int(espera)} segundos...'
                            time.sleep(espera)
                        else:
                            time.sleep(60) # Fallback
                        intentos += 1
                    else:
                        raise e # Si es otro error, que falle
                        
            # Limpiar pedazo
            os.remove(pedazo)

        # 3. RESUMEN FINAL
        TAREAS[task_id]['estado'] = 'Generando Resumen Ejecutivo (Llama 3)...'
        texto_unido = "\n\n".join([f"--- {t['archivo']} ---\n{t['texto']}" for t in transcripciones_list])
        resumen_inteligente = generar_resumen_ia(texto_unido)

        total_menciones = sum(conteo_acumulado.values())
        porcentajes = {cat: (round((val / total_menciones) * 100, 1) if total_menciones > 0 else 0) for cat, val in conteo_acumulado.items()}

        RESULTS_CACHE[task_id] = {
            "resumen": resumen_inteligente, "transcripciones": transcripciones_list,
            "conteo": conteo_acumulado, "porcentajes": porcentajes
        }
        
        TAREAS[task_id]['completado'] = True
        TAREAS[task_id]['estado'] = '¡Análisis Completado!'
        
    except Exception as e:
        TAREAS[task_id]['error'] = str(e)

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/analizar', methods=['POST'])
def analizar_async():
    archivos = request.files.getlist('audios')
    if not archivos or archivos[0].filename == '':
        return jsonify({"error": "No seleccionaste archivos"}), 400
    
    task_id = str(uuid.uuid4())
    TAREAS[task_id] = {'completado': False, 'estado': 'Iniciando carga...', 'error': None}
    
    rutas_guardadas = []
    for archivo in archivos:
        ruta = os.path.join(UPLOAD_FOLDER, f"{uuid.uuid4()}_{archivo.filename}")
        archivo.save(ruta)
        rutas_guardadas.append(ruta)
    
    # Iniciar el trabajo en segundo plano
    thread = threading.Thread(target=procesar_audio_pesado, args=(task_id, rutas_guardadas))
    thread.daemon = True
    thread.start()
    
    return jsonify({"task_id": task_id})

@app.route('/status/<task_id>')
def status(task_id):
    tarea = TAREAS.get(task_id)
    if not tarea: return jsonify({"error": "Tarea no encontrada"}), 404
    
    if tarea.get('error'): return jsonify({"error": tarea['error']})
    if not tarea['completado']: return jsonify({"completado": False, "estado": tarea['estado']})
    
    datos = RESULTS_CACHE[task_id]
    return jsonify({"completado": True, "analysis_id": task_id, "resumen": datos['resumen'], "porcentajes": datos['porcentajes']})

# --- RUTAS DE DESCARGA IGUAL QUE ANTES ---
@app.route('/download/word/<analysis_id>')
def download_word(analysis_id):
    data = RESULTS_CACHE.get(analysis_id)
    if not data: return "Datos expirados.", 404
    document = Document()
    document.add_heading('Reporte de Análisis HELIOS', 0)
    document.add_heading('1. Resumen Ejecutivo', level=1)
    document.add_paragraph(data['resumen'])
    document.add_heading('2. Métricas', level=1)
    for cat, count in data['conteo'].items():
        document.add_paragraph(f"{cat}: {count} menciones ({data['porcentajes'][cat]}%)")
    document.add_heading('3. Transcripciones', level=1)
    for item in data['transcripciones']:
        document.add_paragraph(f"Archivo: {item['archivo']}").bold = True
        document.add_paragraph(item['texto'])
    f = io.BytesIO()
    document.save(f)
    f.seek(0)
    return send_file(f, as_attachment=True, download_name=f'Reporte_Helios_{analysis_id[:8]}.docx')

@app.route('/download/excel/<analysis_id>')
def download_excel(analysis_id):
    data = RESULTS_CACHE.get(analysis_id)
    if not data: return "Datos expirados.", 404
    df_resumen = pd.DataFrame({"Resumen": [data['resumen']]})
    df_metricas = pd.DataFrame(list(data['conteo'].items()), columns=['Categoría', 'Menciones'])
    df_trans = pd.DataFrame(data['transcripciones'])
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df_metricas.to_excel(writer, sheet_name='Métricas', index=False)
        df_resumen.to_excel(writer, sheet_name='Resumen', index=False)
        df_trans.to_excel(writer, sheet_name='Textos', index=False)
    output.seek(0)
    return send_file(output, as_attachment=True, download_name=f'Data_Helios_{analysis_id[:8]}.xlsx')

if __name__ == '__main__':
    app.run(debug=True)